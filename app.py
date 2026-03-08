import streamlit as st
from groq import Groq
import PyPDF2
import docx
import io
import os
from dotenv import load_dotenv



# ─────────────────────────────────────────
#  CLIENT SETUP
# ─────────────────────────────────────────
load_dotenv()
client = Groq(api_key=os.getenv("GROQ_API_KEY"))
# ─────────────────────────────────────────
#  PAGE CONFIG
# ─────────────────────────────────────────
st.set_page_config(
    page_title="ATS Resume Analyzer",
    page_icon="📄",
    layout="centered"
)

# ─────────────────────────────────────────
#  GREEN TECH STYLING
# ─────────────────────────────────────────
st.markdown("""
<style>
    /* Main background */
    .stApp, .main, .block-container,
    section[data-testid="stSidebar"],
    .stFileUploader, div[data-testid="stFileUploadDropzone"],
    div[data-testid="stFileDropzoneInstructions"],
    div[data-baseweb="file-uploader"],
    .stTextArea textarea, .stTextInput input,
    div[data-testid="stAppViewContainer"],
    div[data-testid="stHeader"] {
        background-color: #0d1117 !important;
    }

    /* File uploader box */
    div[data-testid="stFileUploadDropzone"] {
        background-color: #0d1117 !important;
        border: 1px dashed #00ff88 !important;
        border-radius: 8px !important;
    }

    /* Browse files button */
    div[data-testid="stFileUploadDropzone"] button {
        background-color: #00ff88 !important;
        color: #0d1117 !important;
        border-radius: 6px !important;
        font-weight: bold !important;
    }

    /* Title */
    h1 {
        color: #00ff88;
        font-family: 'Courier New', monospace;
        text-align: center;
        text-shadow: 0 0 20px #00ff88;
    }

    /* Subheaders */
    h2, h3 {
        color: #00cc66;
        font-family: 'Courier New', monospace;
    }

    /* Normal text */
    p, label, .stMarkdown, span, div {
        color: #c9d1d9 !important;
    }

    /* Input boxes */
    .stTextArea textarea, .stTextInput input {
        background-color: #0d1117 !important;
        color: #00ff88 !important;
        border: 1px solid #00ff88 !important;
        border-radius: 8px !important;
    }

    /* Analyze button */
    .stButton > button {
        background-color: #00ff88 !important;
        color: #0d1117 !important;
        font-weight: bold !important;
        border-radius: 8px !important;
        font-size: 16px !important;
        width: 100% !important;
        padding: 0.6rem !important;
        border: none !important;
        font-family: 'Courier New', monospace !important;
    }
    .stButton > button:hover {
        background-color: #00cc66 !important;
        box-shadow: 0 0 15px #00ff88 !important;
    }

    /* Result box */
    .result-box {
        background-color: #0d1117;
        border: 1px solid #00ff88;
        border-radius: 10px;
        padding: 1.5rem;
        margin-top: 1rem;
        box-shadow: 0 0 20px rgba(0, 255, 136, 0.1);
    }

    /* Divider */
    hr { border-color: #00ff88 !important; opacity: 0.3; }

    /* Success messages */
    .stSuccess {
        background-color: #0d1117 !important;
        border: 1px solid #00ff88 !important;
        color: #00ff88 !important;
    }

    /* Error messages */
    .stError { border: 1px solid #ff4444 !important; }
</style>
""", unsafe_allow_html=True)

# ─────────────────────────────────────────
#  TITLE
# ─────────────────────────────────────────
st.title("📄 ATS Resume Analyzer")
st.markdown("Upload your resume and paste a job description to get an **ATS score + feedback** and an **improved resume** powered by Groq AI.")
st.divider()


# ─────────────────────────────────────────
#  FUNCTION 1: Extract text from resume
# ─────────────────────────────────────────
def extract_text(uploaded_file):
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text
    elif name.endswith(".docx"):
        doc = docx.Document(io.BytesIO(uploaded_file.read()))
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        return None


# ─────────────────────────────────────────
#  FUNCTION 2: Analyze resume with Groq
# ─────────────────────────────────────────
def analyze_resume(resume_text, job_description):
    prompt = f"""
You are an expert ATS (Applicant Tracking System) analyzer.

Analyze the resume below against the job description and give response in this format:

## ATS Score: [X/100]

## Summary
[2-3 sentence overall summary]

## Strengths
- [strength 1]
- [strength 2]
- [strength 3]

## Missing Keywords
- [missing keyword 1]
- [missing keyword 2]
- [missing keyword 3]

## Suggestions to Improve
- [suggestion 1]
- [suggestion 2]
- [suggestion 3]

### RESUME:
{resume_text}

### JOB DESCRIPTION:
{job_description}
"""
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content


# ─────────────────────────────────────────
#  FUNCTION 3: Generate improved resume
# ─────────────────────────────────────────
def generate_resume(resume_text, analysis_result, job_description):
    prompt = f"""
Based on the original resume and the ATS analysis report below, create an improved
and optimized resume that addresses all the weaknesses and includes the missing keywords.

Format the resume professionally with these sections:
- Full Name & Contact Info (keep from original)
- Professional Summary (improved)
- Skills (add missing keywords)
- Work Experience (improved bullet points)
- Education
- Projects (if any)

Make it ATS friendly and tailored to the job description.

ORIGINAL RESUME:
{resume_text}

ATS ANALYSIS REPORT:
{analysis_result}

JOB DESCRIPTION:
{job_description}
"""
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content


# ─────────────────────────────────────────
#  FUNCTION 4: Save resume as DOCX
# ─────────────────────────────────────────
def save_as_docx(resume_content):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    for line in resume_content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            heading = doc.add_heading(line.replace("# ", ""), level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line.replace("- ", ""), style="List Bullet")
        else:
            doc.add_paragraph(line)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ─────────────────────────────────────────
#  UI — INPUTS
# ─────────────────────────────────────────
st.subheader("📎 Step 1: Upload your Resume")
uploaded_file = st.file_uploader(
    "Upload your resume (PDF or DOCX)",
    type=["pdf", "docx"]
)
st.divider()

st.subheader("📋 Step 2: Paste the Job Description")
job_description = st.text_area(
    "Paste the job description here",
    height=200,
    placeholder="e.g. We are looking for a Python Developer with 2+ years experience..."
)
st.divider()

# ─────────────────────────────────────────
#  UI — ANALYZE BUTTON
# ─────────────────────────────────────────
if st.button("🚀 Analyze My Resume"):
    if not uploaded_file:
        st.error("❌ Please upload your resume.")
    elif not job_description.strip():
        st.error("❌ Please paste a job description.")
    else:
        with st.spinner("🤖 Analyzing your resume..."):
            resume_text = extract_text(uploaded_file)
            if not resume_text or len(resume_text.strip()) < 50:
                st.error("❌ Could not read your resume. Make sure it is not a scanned image.")
            else:
                result = analyze_resume(resume_text, job_description)
                st.success("✅ Analysis Complete!")
                st.markdown('<div class="result-box">', unsafe_allow_html=True)
                st.markdown(result)
                st.markdown('</div>', unsafe_allow_html=True)

                st.divider()

                with st.spinner("📝 Generating your improved resume..."):
                    new_resume = generate_resume(resume_text, result, job_description)

                st.success("✅ Improved Resume Generated!")
                st.markdown('<div class="result-box">', unsafe_allow_html=True)
                st.markdown(new_resume)
                st.markdown('</div>', unsafe_allow_html=True)

                docx_file = save_as_docx(new_resume)
                st.download_button(
                    label="📥 Download Improved Resume as .docx",
                    data=docx_file,
                    file_name="improved_resume.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )